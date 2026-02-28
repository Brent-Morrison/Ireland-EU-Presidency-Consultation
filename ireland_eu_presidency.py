import os
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
import json
from io import BytesIO
from docx import Document
import re
import pandas as pd

# https://github.com/MaartenGr/BERTopic
# https://github.com/laurencleek/thesis_replication/tree/main

BASE_URL = "https://www.gov.ie/en/department-of-foreign-affairs/consultations/irelands-2026-presidency-of-council-of-eu/"


#OUTPUT_DIR = "section7_submissions"
#os.makedirs(OUTPUT_DIR, exist_ok=True)


def get_soup(url):
    r = requests.get(url)
    r.raise_for_status()
    return BeautifulSoup(r.text, "html.parser")


def find_submission_pages():
    """
    Find links to pages listing submissions (e.g. Individual Submissions, Organisation Submissions).
    """
    soup = get_soup(BASE_URL)
    links = []
    for a in soup.find_all("a", href=True):
        text = a.get_text(strip=True).lower()
        if "submissions" in text:
            full_url = urljoin(BASE_URL, a["href"])
            if full_url not in links:
                links.append(full_url)
    return links


def find_individual_submission_links(index_url):
  """
  From a submissions index page, collect links to each individual submission page.
  """
  soup = get_soup(index_url)
  submission_links = []
  for a in soup.find_all("a", href=True):
      href = a["href"]
      # Submission pages are either word or pdf.
      if ".docx" in href or ".pdf" in href:
          full_url = urljoin(index_url, href)
          submission_links.append(full_url)
  return list(set(submission_links))



# Find all submission pages

submission_index_pages = find_submission_pages()

all_submission_links = []
for index_page in submission_index_pages:
    links = find_individual_submission_links(index_page)
    all_submission_links.extend(links)

all_submission_links = list(set(all_submission_links))
print(f"Found {len(all_submission_links)} submission pages.")



# Select test cases
#next(i for i, v in enumerate(all_submission_links) if "_48.docx" in v)
#all_submission_links[92,438,441]
test_links = [all_submission_links[x] for x in [92,192,438,441]]
test_links



def load_doc_from_url(url):
    r = requests.get(url)
    r.raise_for_status()
    return Document(BytesIO(r.content))



def load_doc_network(file_name):
    # 1. Read the local file into memory as bytes
    with open(file_name, 'rb') as f:
        file_content_bytes = f.read()

    # 2. Wrap the bytes in a BytesIO object
    doc_in_memory = BytesIO(file_content_bytes)

    # 3. Now you can use doc_in_memory with docx.Document()
    return Document(doc_in_memory)



def extract_submission(document):
    """
    Extract structured submission data from a consultation DOCX.

    Handles:
    - Mandatory table
    - Optional table
    - Questions 1–5 across split tables
    - Continuation headers
    - Inline responses after instruction phrase
    - Multi-row responses
    """

    result = {}

    # ---------------------------------------------------
    # 1. Extract Mandatory + Optional (first 2 tables)
    # ---------------------------------------------------
    if len(document.tables) < 2:
        raise ValueError("Document does not contain expected tables")

    # ----- Table 1: Mandatory -----
    for row in document.tables[0].rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 2 and cells[0]:
            result[cells[0]] = cells[1]

    # ----- Table 2: Optional -----
    for row in document.tables[1].rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 2 and cells[0]:
            result[cells[0]] = cells[1]

    # ---------------------------------------------------
    # 2. Collect ALL 1-column question tables (after table 2)
    # ---------------------------------------------------
    question_tables = [
        t for t in document.tables[2:]
        if len(t.columns) == 1
    ]

    all_rows = []

    for table in question_tables:
        for row in table.rows:
            text = row.cells[0].text.strip()
            if text:
                all_rows.append(text)

    # ---------------------------------------------------
    # 3. Remove continuation headers and noise
    # ---------------------------------------------------
    HEADER_TEXT1 = "Guiding Questions for Stakeholder Consultations"
    HEADER_TEXT2 = "Through these consultations the Government "

    cleaned_rows = [
        r for r in all_rows
        if HEADER_TEXT1.lower() not in r.lower() or HEADER_TEXT2.lower() not in r.lower()
    ]

    # Remove first two preamble rows if present
    #if len(cleaned_rows) >= 2:
    #    cleaned_rows = cleaned_rows[2:]

    # ---------------------------------------------------
    # 4. Parse Questions + Responses
    # ---------------------------------------------------
    result_section = {}
    current_question = None
    instruction_phrase = "maximum of 500 words."

    for text in cleaned_rows:

        # ---- Case 1: New Question Row ----
        if re.match(r"^Question\s+\d+", text):

            # Extract question label (e.g., "Question 1")
            q_key = text.split("–")[0].strip()

            current_question = q_key
            if current_question not in result_section:
                result_section[current_question] = {"Response": ""}

            # ---- Handle inline response after instruction phrase ----
            if instruction_phrase in text:
                after_instruction = text.split(instruction_phrase, 1)[1].strip()
                if after_instruction:
                    result_section[current_question]["Response"] += (
                        after_instruction + " "
                    )

            continue

        # ---- Case 2: Continuation of current response ----
        if current_question:
            result_section[current_question]["Response"] += text + " "

    # ---------------------------------------------------
    # 5. Final cleanup (strip trailing spaces)
    # ---------------------------------------------------
    for q in result_section:
        result_section[q]["Response"] = result_section[q]["Response"].strip()

    result.update(result_section)

    return result



def transform_for_nlp(extracted_json):
    """
    Convert nested submission JSON into long-format DataFrame
    suitable for NLP analysis.
    """

    rows = []

    for submission_url, submission_data in extracted_json.items():

        # Extract metadata safely
        name = submission_data.get("Name", "")
        organisation = submission_data.get("Organisation (if any)", "")
        respondent_type = submission_data.get(
            "Respondent type (i.e. individual, NGO, business, academic, local authority, etc.)",
            ""
        )
        sector = submission_data.get("What is your sector/area of work?", "")
        geography = submission_data.get(
            "Describe your geographical focus in the context of your submission? For example, rural, urban, national or EU wide.",
            ""
        )

        # Extract submission ID from URL
        submission_id_match = re.search(r"Submission_(\d+)", submission_url)
        submission_id = submission_id_match.group(1) if submission_id_match else submission_url

        # Iterate through question responses
        for key, value in submission_data.items():

            if re.match(r"^Question\s+\d+", key):

                # Handle both dict and string formats
                if isinstance(value, dict):
                    response_text = value.get("Response", "")
                elif isinstance(value, str):
                    response_text = value
                else:
                    continue  # skip unexpected types

                # Skip empty responses
                if not response_text:
                    continue

                question_number = re.search(r"\d+", key).group()

                rows.append({
                    "submission_id": submission_id,
                    "submission_url": submission_url,
                    "organisation": organisation,
                    "respondent_type": respondent_type,
                    "sector": sector,
                    "geography": geography,
                    "question_number": int(question_number),
                    "response_text": response_text
                })

    df = pd.DataFrame(rows)

    def clean_text(text):
        remove_phrases = [
            "Guiding Questions for Stakeholder Consultations",
            "Please limit response to a maximum of 500 words."
        ]

        for phrase in remove_phrases:
            text = text.replace(phrase, "")

        return " ".join(text.split())

    df["response_text"] = df["response_text"].apply(clean_text)

    return df


#url = "https://assets.gov.ie/static/documents/5c48690a/Individual_submission_-_Submission_462.docx"
#doc = load_doc_from_url(url) 
#file_name = "The_Boeing_Company_-_Submission_48.docx"
#file_name = "Individual_submission_-_Submission_462.docx"
#file_name = "Baháí_International_Community_-_Submission_456.docx"
#doc = load_doc_network(file_name)

def main():
    def process_urls(url_list):
        results = {}

        for url in url_list:
            try:
                print(f"Processing {url}")
                doc = load_doc_from_url(url)
                results[url] = extract_submission(doc)
            except Exception as e:
                results[url] = {"error": str(e)}

        return results

    output = process_urls(all_submission_links)
    #print(json.dumps(output, indent=2, ensure_ascii=False))

    df = transform_for_nlp(output)

    df.to_csv("ireland_eu_presidency_submissions.csv", index=False)

if __name__ == "__main__":
    main()